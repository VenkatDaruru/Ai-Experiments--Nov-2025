"""
MULTI-FORMAT DOCUMENT ANALYZER - With Retry Logic
==================================================
"""

import google.generativeai as genai
import os
from dotenv import load_dotenv
from pathlib import Path
from datetime import datetime
import time

# Load API key
load_dotenv()
genai.configure(api_key=os.environ['GEMINI_API_KEY'])

# Use Gemini 2.0 Flash (more generous limits)
model = genai.GenerativeModel('models/gemini-2.0-flash')

def extract_text_from_txt(file_path):
    """Extract text from .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def extract_text_from_docx(file_path):
    """Extract text from .docx file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except ImportError:
        return "ERROR: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"ERROR reading .docx file: {str(e)}"

def extract_text_from_xlsx(file_path):
    """Extract text from .xlsx file"""
    try:
        import pandas as pd
        
        all_sheets = pd.read_excel(file_path, sheet_name=None)
        full_text = []
        
        for sheet_name, df in all_sheets.items():
            full_text.append(f"\n=== SHEET: {sheet_name} ===\n")
            full_text.append(df.to_string())
            
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                full_text.append(f"\n--- Summary Statistics for {sheet_name} ---")
                full_text.append(df[numeric_cols].describe().to_string())
        
        return '\n'.join(full_text)
    
    except ImportError:
        return "ERROR: pandas and openpyxl not installed. Run: pip install pandas openpyxl"
    except Exception as e:
        return f"ERROR reading .xlsx file: {str(e)}"

def analyze_with_retry(prompt, max_retries=3):
    """Analyze with automatic retry on rate limits"""
    
    for attempt in range(max_retries):
        try:
            print(f"\n🤖 Analyzing with Gemini AI (Attempt {attempt + 1}/{max_retries})...")
            response = model.generate_content(prompt)
            
            if not response.candidates:
                print("❌ Analysis blocked for safety reasons")
                return None
            
            # Success!
            print(f"✓ Analysis complete!")
            print(f"📊 Tokens used: {response.usage_metadata.total_token_count}")
            cost = (response.usage_metadata.total_token_count / 1_000_000) * 0.15
            print(f"💰 Cost: ${cost:.6f}")
            
            return response.text
        
        except Exception as e:
            error_str = str(e)
            
            # Check if it's a rate limit error
            if "429" in error_str or "quota" in error_str.lower():
                wait_time = 60 * (attempt + 1)  # Exponential backoff
                print(f"\n⚠️ Rate limit hit! Waiting {wait_time} seconds...")
                print(f"   (This is normal - just a speed limit, not a billing issue)")
                
                if attempt < max_retries - 1:
                    for remaining in range(wait_time, 0, -5):
                        print(f"   Retrying in {remaining} seconds...", end='\r')
                        time.sleep(5)
                    print()  # New line
                else:
                    print(f"\n❌ Still hitting rate limits after {max_retries} attempts")
                    print("   Suggestions:")
                    print("   1. Wait 2-3 minutes and try again")
                    print("   2. Try a smaller document")
                    print("   3. Check your usage at: https://ai.dev/usage")
                    return None
            else:
                print(f"\n❌ Error during analysis: {error_str}")
                return None
    
    return None

def analyze_document(file_path):
    """Analyze document using appropriate extractor"""
    
    print(f"\n{'='*60}")
    print(f"ANALYZING: {Path(file_path).name}")
    print(f"{'='*60}\n")
    
    if not os.path.exists(file_path):
        print(f"❌ ERROR: File not found: {file_path}")
        return None
    
    extension = Path(file_path).suffix.lower()
    print(f"📄 File type: {extension}")
    
    # Extract text based on file type
    if extension == '.txt':
        print("📖 Extracting text from .txt file...")
        document_text = extract_text_from_txt(file_path)
    
    elif extension == '.docx':
        print("📖 Extracting text from .docx file...")
        document_text = extract_text_from_docx(file_path)
    
    elif extension in ['.xlsx', '.xls']:
        print("📖 Extracting data from .xlsx file...")
        document_text = extract_text_from_xlsx(file_path)
    
    else:
        print(f"❌ ERROR: Unsupported file type: {extension}")
        print(f"Supported types: .txt, .docx, .xlsx")
        return None
    
    if document_text.startswith("ERROR"):
        print(f"\n❌ {document_text}")
        return None
    
    if not document_text.strip():
        print("❌ ERROR: File appears to be empty")
        return None
    
    print(f"✓ Extracted {len(document_text)} characters")
    
    # Limit text size to avoid huge requests
    max_chars = 50000  # About 12,500 tokens
    if len(document_text) > max_chars:
        print(f"\n⚠️ Document is large ({len(document_text)} chars)")
        print(f"   Truncating to first {max_chars} characters to avoid rate limits")
        document_text = document_text[:max_chars] + "\n\n[Document truncated...]"
    
    prompt = f"""
    Analyze the following document and extract:
    
    1. DOCUMENT TYPE: What kind of document is this?
    
    2. SUMMARY: A brief 2-3 sentence summary
    
    3. KEY POINTS: Main topics (bullet points)
    
    4. ACTION ITEMS: Any tasks or actions (with owners if specified)
    
    5. IMPORTANT DATES: Any dates or deadlines
    
    6. RISKS/CONCERNS: Any issues or problems
    
    7. NUMBERS/METRICS: Important statistics or data
    
    Format your response clearly with these headers.
    
    DOCUMENT CONTENT:
    {document_text}
    """
    
    # Use retry logic
    return analyze_with_retry(prompt)

def save_analysis(analysis, original_file):
    """Save analysis to a file"""
    if not analysis:
        return None
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    original_name = Path(original_file).stem
    output_file = f"analysis_{original_name}_{timestamp}.txt"
    
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"ANALYSIS OF: {original_file}\n")
            f.write(f"GENERATED: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*60 + "\n\n")
            f.write(analysis)
        
        print(f"\n💾 Analysis saved to: {output_file}")
        return output_file
        
    except Exception as e:
        print(f"\n❌ Error saving file: {str(e)}")
        return None

def main():
    """Main function"""
    print("="*60)
    print("MULTI-FORMAT DOCUMENT ANALYZER v2.0")
    print("With Automatic Retry & Rate Limit Handling")
    print("="*60)
    print("\nSupported formats: .txt, .docx, .xlsx")
    print("Using: Gemini 2.0 Flash (generous rate limits)")
    
    file_path = input("\nEnter the path to your document: ").strip()
    file_path = file_path.strip('"').strip("'")
    
    if not file_path:
        print("\n❌ No file specified. Exiting.")
        return
    
    analysis = analyze_document(file_path)
    
    if analysis:
        print("\n" + "="*60)
        print("ANALYSIS RESULTS")
        print("="*60 + "\n")
        print(analysis)
        
        save_analysis(analysis, file_path)
        
        print("\n" + "="*60)
        print("✅ Analysis complete!")
        print("="*60)
    else:
        print("\n❌ Analysis failed. Please check the errors above.")

if __name__ == "__main__":
    main()